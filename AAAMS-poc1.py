import streamlit as st
import pandas as pd
import plotly.graph_objects as go
from openai import OpenAI
import numpy as np
import io
from docx import Document
from docx.shared import Inches
import tempfile

# Set page title
st.set_page_config(page_title="Excel Data Processor", layout="wide")

# Sidebar for configuration
st.sidebar.header("Configuration")

# Password protection
password = st.sidebar.text_input("Enter password", type="password")
correct_password = "iken"  # Replace with your desired password

# Initialize session state
if 'data_loaded' not in st.session_state:
    st.session_state.data_loaded = False
if 'new_df' not in st.session_state:
    st.session_state.new_df = None
if 'selected_company' not in st.session_state:
    st.session_state.selected_company = None
if 'ratings_generated' not in st.session_state:
    st.session_state.ratings_generated = False
if 'edits_confirmed' not in st.session_state:
    st.session_state.edits_confirmed = False
if 'executive_summary' not in st.session_state:
    st.session_state.executive_summary = ""

# Hardcoded template path
template_path = "template.docx"

def split_header(header):
    parts = header.split("-", 1)
    if len(parts) == 2:
        return parts[0].rstrip(), parts[1].lstrip()
    return "", header

def create_spider_chart(df):
    categories = df['Kategorie'].unique()
    avg_ratings = df.groupby('Kategorie')['Bewertung (1-5)'].mean().values

    fig = go.Figure(data=[
        go.Scatterpolar(
            r=avg_ratings,
            theta=categories,
            fill='toself',
            line=dict(color='rgb(31, 119, 180)'),  # Blue color
        )
    ])

    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,
                range=[0, 5]
            )
        ),
        showlegend=False,
        paper_bgcolor='rgba(0,0,0,0)',  # Transparent background
        plot_bgcolor='rgba(0,0,0,0)'    # Transparent plot area
    )

    return fig

def generate_executive_summary(df):
    client = OpenAI(api_key=st.secrets["openai_api_key"])

    # Prepare the data for the prompt
    categories = df.groupby('Kategorie')['Bewertung (1-5)'].mean().reset_index()
    categories['Bewertung (1-5)'] = categories['Bewertung (1-5)'].round(2)
    
    prompt = f"""Erstelle eine Executive Summary basierend auf den folgenden Daten:

Gesamtbewertungen nach Kategorie:
{categories.to_string(index=False)}

Detaillierte Informationen:
{df.to_string(index=False)}

Bitte erstelle eine umfassende Executive Summary, die:
1. Die durchschnittlichen Bewertungen für jede Kategorie hervorhebt
2. Bemerkenswerte Einzelaspekte aus jeder Kategorie diskutiert
3. Eine Gesamtbewertung der Leistung des Unternehmens über alle Kategorien hinweg liefert

Die Zusammenfassung sollte prägnant, aber informativ sein und für eine Überprüfung auf Führungsebene geeignet sein. 
Bitte verwende Schweizer Rechtschreibung und Grammatik. Starte den Bericht nicht mit einem Titel, sonder beginne direkt mit dem Inhalt."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.7
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Fehler bei der Generierung der Executive Summary: {e}"

def generate_document(spider_chart, df):
    try:
        # Load the Word template
        template = Document(template_path)

        # Save the spider chart to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmpfile:
            spider_chart.write_image(tmpfile.name, format="png", scale=2)  # Higher resolution
            
            # Replace placeholders in the template
            for paragraph in template.paragraphs:
                if "{{Executive_Summary}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{Executive_Summary}}", st.session_state.executive_summary)
                if "{{Company_Name}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{Company_Name}}", st.session_state.selected_company)
                if "{{Spider_Chart}}" in paragraph.text:
                    # Replace the placeholder with the image
                    run = paragraph.runs[0]
                    run.text = ""
                    run.add_picture(tmpfile.name, width=Inches(6))

        # Testing block: Print all placeholders and their replacements
        print("Placeholder Replacements:")
        for table in template.tables:
            for row_index, row in enumerate(table.rows):
                for col_index, cell in enumerate(row.cells):
                    placeholder = f"{{{{{'abcd'[col_index]}{row_index}}}}}"
                    if placeholder in cell.text:
                        if row_index == 0:
                            # Header row
                            replacement = df.columns[col_index]
                        elif row_index - 1 < len(df):
                            # Data rows
                            replacement = str(df.iloc[row_index - 1, col_index])
                        else:
                            replacement = "N/A"  # For rows beyond DataFrame length
                        print(f"Placeholder: {placeholder}, Replacement: {replacement}")

        # Replace placeholders in the table
        for table in template.tables:
            for row_index, row in enumerate(table.rows):
                for col_index, cell in enumerate(row.cells):
                    placeholder = f"{{{{{'abcd'[col_index]}{row_index}}}}}"
                    if placeholder in cell.text:
                        if row_index == 0:
                            # Header row
                            replacement = df.columns[col_index]
                        elif row_index - 1 < len(df):
                            # Data rows
                            replacement = str(df.iloc[row_index - 1, col_index])
                        else:
                            replacement = "N/A"  # For rows beyond DataFrame length
                        cell.text = cell.text.replace(placeholder, replacement)

        # Save the generated document to a buffer
        buffer = io.BytesIO()
        template.save(buffer)
        buffer.seek(0)

        return buffer
    except Exception as e:
        st.error(f"Fehler bei der Generierung des Word-Dokuments: {e}")
        return None

if password == correct_password:
    # File uploader
    uploaded_file = st.sidebar.file_uploader("Upload your Excel file", type=["xlsx"])

    if uploaded_file is not None and not st.session_state.data_loaded:
        # Read the Excel file
        df = pd.read_excel(uploaded_file, sheet_name="Sheet1", header=0)

        # Create a dropdown with company names from column F (index 5)
        company_names = df.iloc[:, 5].dropna().unique()
        selected_company = st.sidebar.selectbox("Select a company", company_names)

        if st.sidebar.button("Load Data"):
            # Store the selected company name in session state
            st.session_state.selected_company = selected_company

            # Filter the dataframe for the selected company
            selected_row = df[df.iloc[:, 5] == selected_company].iloc[0]

            # Process the headers and create the new dataframe
            headers = df.columns[7:]  # Start from column H (index 7)
            answers = selected_row.iloc[7:]  # Corresponding answers

            # Split headers into Kategorie and Frage
            kategorien, fragen = zip(*[split_header(header) for header in headers])

            # Create a new dataframe with Kategorie, Frage, and Answer
            new_df = pd.DataFrame({
                "Kategorie": kategorien,
                "Frage": fragen,
                "Antwort": answers,
                "Bewertung (1-5)": [""] * len(answers)  # Empty column for ratings
            })

            # Store the new dataframe in session state
            st.session_state.new_df = new_df
            st.session_state.data_loaded = True
            st.session_state.ratings_generated = False
            st.session_state.edits_confirmed = False

    if st.session_state.data_loaded and st.session_state.new_df is not None:
        st.header(f"Data for {st.session_state.selected_company}")
        
        # Create a placeholder for the dataframe
        table_placeholder = st.empty()
        
        if not st.session_state.ratings_generated:
            # Display the initial dataframe
            table_placeholder.dataframe(st.session_state.new_df, use_container_width=True, hide_index=True)

            if st.button("Generate Ratings"):
                # Initialize the OpenAI client
                client = OpenAI(api_key=st.secrets["openai_api_key"])

                # Generate ratings using GPT API
                ratings = []
                progress_bar = st.progress(0)
                total_rows = len(st.session_state.new_df)
                
                for counter, (_, row) in enumerate(st.session_state.new_df.iterrows(), 1):
                    prompt = f"""Erstelle mir eine Bewertung des folgenden Textes, nach einem Punkte System 1-5 (1 schlecht und 5 gut):
Kategorie: {row['Kategorie']}
Frage: {row['Frage']}
Antwort: {row['Antwort']}
Gib nur eine ganze Zahl zwischen 1 und 5 zurück. Wenn keine Antwort vorhanden ist oder die Daten für eine Bewertung unzureichend sind, gib 1 zurück."""
                    try:
                        response = client.chat.completions.create(
                            model="gpt-4o",
                            messages=[
                                {
                                    "role": "user",
                                    "content": prompt
                                }
                            ],
                            temperature=0.7
                        )
                        rating = response.choices[0].message.content.strip()
                        # Ensure the rating is an integer between 1 and 5
                        rating = max(1, min(5, int(float(rating))))
                        ratings.append(rating)
                    except Exception as e:
                        st.error(f"Fehler bei der Generierung der Bewertung: {e}")
                        ratings.append(1)  # Default to 1 in case of an error
                    
                    # Update progress bar
                    progress_bar.progress(counter / total_rows)

                # Add ratings to the dataframe
                st.session_state.new_df["Bewertung (1-5)"] = ratings

                # Remove the progress bar
                progress_bar.empty()

                st.session_state.ratings_generated = True

        if st.session_state.ratings_generated and not st.session_state.edits_confirmed:
            # Allow editing of ratings
            edited_df = st.data_editor(
                st.session_state.new_df,
                column_config={
                    "Bewertung (1-5)": st.column_config.NumberColumn(
                        "Bewertung (1-5)",
                        min_value=1,
                        max_value=5,
                        step=1,
                        format="%d"
                    )
                },
                use_container_width=True,
                hide_index=True,
                disabled=["Kategorie", "Frage", "Antwort"],
            )

            # Update the session state with the edited dataframe
            st.session_state.new_df = edited_df

            if st.button("Confirm Edits"):
                st.session_state.edits_confirmed = True

        if st.session_state.edits_confirmed:
            # Convert ratings to integers and handle any non-numeric values
            st.session_state.new_df['Bewertung (1-5)'] = pd.to_numeric(st.session_state.new_df['Bewertung (1-5)'], errors='coerce').fillna(1).clip(1, 5).astype(int)

            # Create and display the spider chart
            st.subheader("Category Ratings Overview")
            spider_chart = create_spider_chart(st.session_state.new_df)
            st.plotly_chart(spider_chart, use_container_width=True)

            # Generate and display the executive summary
            if 'executive_summary' not in st.session_state or not st.session_state.executive_summary:
                with st.spinner("Generiere Executive Summary..."):
                    st.session_state.executive_summary = generate_executive_summary(st.session_state.new_df)
            
            st.subheader("Executive Summary")
            st.write(st.session_state.executive_summary)

            # Consolidated button for generation and download
            if st.button("Generate and download Word document"):
                with st.spinner("Generiere Dokument..."):
                    doc_buffer = generate_document(spider_chart, st.session_state.new_df)
                    if doc_buffer:
                        st.download_button(
                            label="Klicken Sie hier, um das generierte Dokument herunterzuladen",
                            data=doc_buffer,
                            file_name="generated_report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_button"
                        )
                        st.success("Dokument erfolgreich generiert! Klicken Sie auf den Download-Button oben, um es zu speichern.")
                    else:
                        st.error("Fehler bei der Generierung des Dokuments. Bitte versuchen Sie es erneut.")

else:
    st.error("Falsches Passwort. Bitte geben Sie das korrekte Passwort ein, um auf die Anwendung zuzugreifen.")
